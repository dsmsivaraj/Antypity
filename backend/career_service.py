from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Dict, List, Optional
from urllib.parse import quote_plus, urlparse
from uuid import uuid4

import httpx
from bs4 import BeautifulSoup

from .config import Settings
from .database import PostgreSQLDatabaseClient
from .model_router import ModelRouter
from .prompt_registry import register_prompt

TRUSTED_JOB_PORTALS: Dict[str, Dict[str, Any]] = {
    "linkedin": {
        "label": "LinkedIn Jobs",
        "hosts": ["linkedin.com", "www.linkedin.com"],
        "search_url": "https://www.linkedin.com/jobs/search/?keywords={keywords}&location={location}",
    },
    "indeed": {
        "label": "Indeed",
        "hosts": ["indeed.com", "www.indeed.com"],
        "search_url": "https://www.indeed.com/jobs?q={keywords}&l={location}",
    },
    "glassdoor": {
        "label": "Glassdoor",
        "hosts": ["glassdoor.com", "www.glassdoor.com"],
        "search_url": "https://www.glassdoor.com/Job/jobs.htm?sc.keyword={keywords}&locT=C&locId=&locKeyword={location}",
    },
    "wellfound": {
        "label": "Wellfound",
        "hosts": ["wellfound.com", "angel.co"],
        "search_url": "https://wellfound.com/jobs?query={keywords}&location={location}",
    },
    "naukri": {
        "label": "Naukri",
        "hosts": ["naukri.com", "www.naukri.com"],
        "search_url": "https://www.naukri.com/{keywords}-jobs-in-{location}",
    },
    "dice": {
        "label": "Dice",
        "hosts": ["dice.com", "www.dice.com"],
        "search_url": "https://www.dice.com/jobs?q={keywords}&location={location}",
    },
    "ziprecruiter": {
        "label": "ZipRecruiter",
        "hosts": ["ziprecruiter.com", "www.ziprecruiter.com"],
        "search_url": "https://www.ziprecruiter.com/jobs-search?search={keywords}&location={location}",
    },
}

DEFAULT_TEMPLATES: List[Dict[str, Any]] = [
    {
        "id": "modern-executive",
        "name": "Modern Executive",
        "target_role": "Senior engineering and product leadership",
        "style": "executive-minimal",
        "figma_prompt": "Create a clean one-page resume with a bold name band, asymmetric metrics rail, and restrained serif/sans pairing.",
        "sections": ["Header", "Executive Summary", "Experience", "Leadership Impact", "Core Skills", "Education"],
        "design_tokens": {
            "primary": "#0d3b66",
            "accent": "#f4a261",
            "surface": "#f8fafc",
            "font_heading": "Manrope",
            "font_body": "IBM Plex Sans",
        },
        "preview_markdown": "A one-page executive resume with a strong summary and quantified impact blocks.",
        "source": "system",
    },
    {
        "id": "creative-product",
        "name": "Creative Product Builder",
        "target_role": "Product, design, and startup operators",
        "style": "editorial-modular",
        "figma_prompt": "Design a modular resume with editorial hierarchy, stacked project cards, and subtle grid accents.",
        "sections": ["Header", "Profile", "Selected Work", "Experience", "Tools", "Education"],
        "design_tokens": {
            "primary": "#1d3557",
            "accent": "#e63946",
            "surface": "#fffaf3",
            "font_heading": "Space Grotesk",
            "font_body": "DM Sans",
        },
        "preview_markdown": "A portfolio-forward resume with project callouts and modular content blocks.",
        "source": "system",
    },
]


@dataclass(frozen=True)
class ParsedResume:
    filename: str
    text: str
    metadata: Dict[str, Any]
    parsed_fields: Dict[str, Any]


class CareerService:
    def __init__(
        self,
        *,
        settings: Settings,
        model_router: ModelRouter,
        database_client: PostgreSQLDatabaseClient,
        embedding_service=None,
    ) -> None:
        self.settings = settings
        self.model_router = model_router
        self.database_client = database_client
        self._embedding_service = embedding_service

    def trusted_sources(self) -> List[Dict[str, Any]]:
        allowed = set(self.settings.trusted_job_sources)
        return [
            {"id": key, **value}
            for key, value in TRUSTED_JOB_PORTALS.items()
            if key in allowed
        ]

    def parse_resume(self, filename: str, content: bytes) -> ParsedResume:
        lower_name = filename.lower()
        if lower_name.endswith(".pdf"):
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(content))
            raw_pages = [(page.extract_text() or "").strip() for page in reader.pages]
            text = self._fix_pdf_text("\n".join(raw_pages)).strip()
            metadata = {"pages": len(reader.pages), "type": "pdf"}
        elif lower_name.endswith(".docx"):
            import docx

            doc = docx.Document(BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()
            metadata = {"paragraphs": len(doc.paragraphs), "type": "docx"}
        elif lower_name.endswith(".txt"):
            text = content.decode("utf-8", errors="ignore").strip()
            metadata = {"type": "txt"}
        else:
            raise ValueError("Unsupported file format. Supported formats: pdf, docx, txt.")

        from .resume_parser import parse_resume_text
        parsed_fields = parse_resume_text(text)
        return ParsedResume(filename=filename, text=text, metadata=metadata, parsed_fields=parsed_fields)

    def analyze_resume(
        self,
        *,
        resume_text: str,
        jd_text: str,
        model_profile: Optional[str],
        source_filename: Optional[str],
        created_by: Optional[str],
    ) -> Dict[str, Any]:
        resume_keywords = self._extract_keywords(resume_text)
        jd_keywords = self._extract_keywords(jd_text)
        matched = sorted(resume_keywords & jd_keywords)
        missing = sorted(jd_keywords - resume_keywords)
        strengths = self._extract_bullets(resume_text, limit=4)
        summary_seed = (
            f"Resume keywords: {', '.join(sorted(resume_keywords)[:20]) or 'none'}\n"
            f"JD keywords: {', '.join(sorted(jd_keywords)[:20]) or 'none'}\n"
            f"Top strengths: {', '.join(strengths) or 'none'}"
        )
        profile, completion = self.model_router.complete(
            model_profile=model_profile or self._preferred_profile("resume"),
            prompt=summary_seed,
            system_prompt=(
                "You are a resume reviewer. Summarize the candidate fit in 4-6 concise sentences, "
                "highlight missing experience, and keep the answer grounded in the provided signals."
            ),
        )
        suggestions = self._build_resume_suggestions(matched=matched, missing=missing, strengths=strengths)
        recommended_roles = self._recommend_roles(resume_keywords, resume_text)
        match_score = self._compute_match_score(matched=len(matched), total=max(len(jd_keywords), 1))
        record = {
            "id": str(uuid4()),
            "title": source_filename or "Uploaded resume",
            "source_filename": source_filename,
            "resume_text": resume_text,
            "jd_text": jd_text,
            "summary": completion.content,
            "match_score": match_score,
            "suggestions": suggestions,
            "ats_keywords": sorted(matched) if jd_text else sorted(resume_keywords)[:25],
            "strengths": strengths,
            "gaps": missing[:12],
            "recommended_roles": recommended_roles,
            "model_profile": profile.id,
            "created_by": created_by,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "used_llm": completion.used_llm,
            "provider": completion.provider,
        }
        try:
            # register prompt version for governance
            pmeta = register_prompt("resume_analyze", summary_seed)
            record.setdefault('metadata', {})
            record['metadata']['prompt_version'] = pmeta.get('version')
        except Exception:
            pass

        # persist analysis
        if self.database_client.is_configured:
            self.database_client.save_resume_analysis(record)

        # create section embeddings for RAG
        try:
            from .embeddings_service import get_embedding_service

            embedding_svc = self._embedding_service or get_embedding_service()

            # split resume into heuristic sections by headings
            sections = []
            lines = resume_text.splitlines()
            current = []
            section_id = 0
            for ln in lines:
                if ln.strip().endswith((':', '：')) or ln.strip().lower() in ('experience', 'education', 'skills', 'summary'):
                    if current:
                        sections.append(('section-' + str(section_id), '\n'.join(current)))
                        section_id += 1
                        current = []
                    current.append(ln)
                else:
                    current.append(ln)
            if current:
                sections.append(('section-' + str(section_id), '\n'.join(current)))

            for sid, stext in sections:
                embedding_svc.index_document(record["id"], sid, stext)
        except Exception:
            # embedding service is optional — ignore failures in environments without it
            pass

        return record

    def chat_resume(
        self,
        *,
        question: str,
        resume_text: str,
        jd_text: str,
        model_profile: Optional[str],
    ) -> Dict[str, Any]:
        contexts: List[Dict[str, Any]] = []
        retrieval_used = False
        # Retrieve relevant resume sections for RAG
        try:
            from .retrieval import get_context_blocks

            contexts = get_context_blocks(question + "\n" + resume_text, top_k=3)
            retrieval_used = len(contexts) > 0
            context_text = "\n\n".join([f"[{c['doc_id']}:{c['section_id']}] {c['text']}" for c in contexts])
            prompt_body = (
                "Relevant resume sections:\n" + context_text + "\n\n"
                f"Resume:\n{resume_text[:5000]}\n\n"
                f"Job description:\n{jd_text[:5000]}\n\n"
                f"Question:\n{question}"
            )
        except Exception:
            # retrieval is optional; fall back to original prompt
            prompt_body = (
                f"Resume:\n{resume_text[:5000]}\n\n"
                f"Job description:\n{jd_text[:5000]}\n\n"
                f"Question:\n{question}"
            )

        profile, completion = self.model_router.complete(
            model_profile=model_profile or self._preferred_profile("resume"),
            prompt=prompt_body,
            system_prompt=(
                "You are a resume copilot. Answer precisely, quote relevant experience, suggest edits if needed, "
                "and keep the reply concise and recruiter-friendly."
            ),
        )
        try:
            pmeta = register_prompt("resume_chat", prompt_body)
            prompt_version = pmeta.get('version')
        except Exception:
            prompt_version = None
        grounding_score = self._score_resume_chat_grounding(contexts)
        confidence = self._confidence_from_grounding(
            grounding_score=grounding_score,
            citation_count=len(contexts[:3]),
            drift_flag=not retrieval_used,
        )
        self._record_quality_metric(
            response_type="resume_chat",
            grounding_score=grounding_score,
            citation_count=len(contexts[:3]),
            confidence=confidence,
            drift_flag=not retrieval_used,
            metadata={
                "question": question[:500],
                "model_profile": profile.id,
                "used_llm": completion.used_llm,
            },
        )
        return {
            "answer": completion.content,
            "used_llm": completion.used_llm,
            "provider": completion.provider,
            "model_profile": profile.id,
            "citations": [
                {
                    "doc_id": c.get("doc_id"),
                    "section_id": c.get("section_id"),
                    "excerpt": c.get("text"),
                    "score": round(float(c.get("score", 0.0)), 4),
                    "source_type": "resume_embedding",
                }
                for c in contexts[:3]
            ],
            "confidence": confidence,
            "prompt_version": prompt_version,
            "suggested_questions": [
                "Which bullets should I rewrite for ATS?",
                "What keywords are still missing for this role?",
                "How should I tighten the summary for this position?",
                "Draft a cover letter for this job.",
                "Where should I look for recruiter or HR contacts for this role?",
            ],
        }

    def create_cover_letter(
        self,
        *,
        resume_text: str,
        jd_text: str,
        target_role: str,
        company_name: str,
        hiring_manager_name: str,
        tone: str,
        model_profile: Optional[str],
        created_by: Optional[str],
    ) -> Dict[str, Any]:
        manager_line = hiring_manager_name.strip() or "Hiring Manager"
        profile, completion = self.model_router.complete(
            model_profile=model_profile or self._preferred_profile("resume"),
            prompt=(
                f"Target role: {target_role}\n"
                f"Company: {company_name}\n"
                f"Hiring manager: {manager_line}\n"
                f"Tone: {tone}\n\n"
                f"Resume:\n{resume_text[:5000]}\n\n"
                f"Job description:\n{jd_text[:4000]}\n"
            ),
            system_prompt=(
                "You are an expert career strategist and recruiter-facing writer. "
                "Write a concise, evidence-based cover letter that sounds credible, avoids cliches, "
                "anchors claims in resume evidence, and mirrors the target role's language. "
                "Start with a compelling subject line on the first line using the format 'Subject: ...', "
                "then write the cover letter body in 3 short paragraphs, and finish with a strong closing sentence. "
                "After the letter, add a section called 'Talking Points:' with 3 bullet points."
            ),
        )
        subject_line, cover_letter, talking_points = self._split_cover_letter_response(completion.content, target_role, company_name)
        record = {
            "id": str(uuid4()),
            "query_type": "cover_letter",
            "query_text": f"{company_name}::{target_role}",
            "sources": ["resume", "job_description"],
            "result_count": len(talking_points),
            "metadata": {
                "company_name": company_name,
                "target_role": target_role,
                "subject_line": subject_line,
                "tone": tone,
                "model_profile": profile.id,
            },
            "created_by": created_by,
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        try:
            prompt_text = (
                f"Target role: {target_role}\n"
                f"Company: {company_name}\n"
                f"Hiring manager: {manager_line}\n"
                f"Tone: {tone}\n\n"
                f"Resume:\n{resume_text[:5000]}\n\n"
                f"Job description:\n{jd_text[:4000]}\n"
            )
            pmeta = register_prompt("cover_letter", prompt_text)
            record.setdefault('metadata', {})
            record['metadata']['prompt_version'] = pmeta.get('version')
        except Exception:
            pass

        if self.database_client.is_configured:
            self.database_client.save_career_query(record)
        evidence_terms = self._shared_evidence_terms(resume_text, jd_text, limit=6)
        citations = self._build_text_citations(
            sources=[
                ("resume", resume_text),
                ("job_description", jd_text),
            ],
            focus_terms=evidence_terms or self._extract_keywords(target_role),
            limit=4,
        )
        grounding_score = self._score_grounding(citations=citations, evidence_terms=evidence_terms, target_terms=6)
        confidence = self._confidence_from_grounding(
            grounding_score=grounding_score,
            citation_count=len(citations),
            drift_flag=not citations,
        )
        self._record_quality_metric(
            response_type="cover_letter",
            grounding_score=grounding_score,
            citation_count=len(citations),
            confidence=confidence,
            drift_flag=not citations,
            metadata={
                "company_name": company_name,
                "target_role": target_role,
                "model_profile": profile.id,
                "used_llm": completion.used_llm,
            },
        )
        return {
            "company_name": company_name,
            "target_role": target_role,
            "subject_line": subject_line,
            "cover_letter": cover_letter,
            "talking_points": talking_points,
            "citations": citations,
            "confidence": confidence,
            "used_llm": completion.used_llm,
            "provider": completion.provider,
            "model_profile": profile.id,
            "prompt_version": record.get('metadata', {}).get('prompt_version'),
        }

    async def discover_recruiter_contacts(
        self,
        *,
        company_name: str,
        company_domain: str,
        job_url: str,
        source_text: str,
        target_role: str,
        created_by: Optional[str],
    ) -> Dict[str, Any]:
        normalized_domain = self._normalize_company_domain(company_domain or self._derive_company_domain(job_url))
        contacts: List[Dict[str, Any]] = []
        lookup_urls = self._build_contact_lookup_urls(company_name, normalized_domain, target_role)

        if source_text.strip():
            contacts.extend(self._extract_contacts_from_text(source_text, source="provided-text"))

        if normalized_domain:
            page_urls = self._company_lookup_pages(normalized_domain)
            async with httpx.AsyncClient(timeout=8.0, headers={"User-Agent": "Actypity/2.1"}) as client:
                for url in page_urls:
                    try:
                        response = await client.get(url)
                        if response.status_code >= 400:
                            continue
                        contacts.extend(self._extract_contacts_from_text(response.text, source=url))
                    except httpx.HTTPError:
                        continue

            if not any(contact.get("email") for contact in contacts):
                contacts.extend(self._inferred_role_mailboxes(company_name, normalized_domain))

        deduped = self._dedupe_contacts(contacts)
        verified_count = sum(1 for contact in deduped if contact.get("confidence") in {"high", "medium"})
        inferred_count = sum(1 for contact in deduped if contact.get("confidence") == "low")
        provenance = self._contact_provenance(deduped, lookup_urls)
        confidence = self._confidence_from_grounding(
            grounding_score=self._contact_grounding_score(verified_count=verified_count, total_count=len(deduped)),
            citation_count=len(provenance),
            drift_flag=verified_count == 0 and inferred_count > 0,
        )
        if self.database_client.is_configured:
            self.database_client.save_career_query(
                {
                    "id": str(uuid4()),
                    "query_type": "recruiter_contacts",
                    "query_text": company_name,
                    "sources": [url for url in lookup_urls[:8]],
                    "result_count": len(deduped),
                    "metadata": {
                        "company_domain": normalized_domain,
                        "job_url": job_url,
                        "target_role": target_role,
                    },
                    "created_by": created_by,
                    "created_at": datetime.now(timezone.utc).isoformat(),
                }
            )
        self._record_quality_metric(
            response_type="recruiter_contacts",
            grounding_score=self._contact_grounding_score(verified_count=verified_count, total_count=len(deduped)),
            citation_count=len(provenance),
            confidence=confidence,
            drift_flag=verified_count == 0 and inferred_count > 0,
            metadata={
                "company_name": company_name,
                "company_domain": normalized_domain,
                "verified_contact_count": verified_count,
                "inferred_contact_count": inferred_count,
            },
        )

        return {
            "company_name": company_name,
            "company_domain": normalized_domain or None,
            "contacts": deduped,
            "lookup_urls": lookup_urls,
            "verified_contact_count": verified_count,
            "inferred_contact_count": inferred_count,
            "confidence": confidence,
            "provenance": provenance,
        }

    def list_templates(self) -> List[Dict[str, Any]]:
        templates = list(DEFAULT_TEMPLATES)
        if self.database_client.is_configured:
            templates.extend(
                [
                    {**record, "source": "generated"}
                    for record in self.database_client.list_resume_templates(limit=50)
                ]
            )
        return templates

    def design_template(
        self,
        *,
        name: str,
        target_role: str,
        style: str,
        notes: str,
        model_profile: Optional[str],
        created_by: Optional[str],
    ) -> Dict[str, Any]:
        sections = self._template_sections_for_role(target_role)
        design_tokens = self._design_tokens_for_style(style)
        profile, completion = self.model_router.complete(
            model_profile=model_profile or self._preferred_profile("template"),
            prompt=(
                f"Template name: {name}\n"
                f"Target role: {target_role}\n"
                f"Style: {style}\n"
                f"Notes: {notes or 'No additional notes.'}\n"
                f"Sections: {', '.join(sections)}"
            ),
            system_prompt=(
                "You are a resume art director. Produce a Figma-ready creative brief with layout direction, "
                "type hierarchy, spacing, and content framing for a high-conversion resume template."
            ),
        )
        record = {
            "id": str(uuid4()),
            "name": name,
            "target_role": target_role,
            "style": style,
            "notes": notes,
            "figma_prompt": completion.content,
            "sections": sections,
            "design_tokens": design_tokens,
            "preview_markdown": self._preview_markdown(name, target_role, sections),
            "model_profile": profile.id,
            "created_by": created_by,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "source": "generated",
        }
        if self.database_client.is_configured:
            self.database_client.save_resume_template(record)
        return record

    async def extract_job_description(
        self,
        *,
        url: Optional[str],
        text: Optional[str],
    ) -> Dict[str, Any]:
        if text and text.strip():
            cleaned = text.strip()
            citations = self._build_text_citations(
                sources=[("manual", cleaned)],
                focus_terms=self._extract_keywords(cleaned),
                limit=3,
            )
            confidence = self._confidence_from_grounding(
                grounding_score=self._score_grounding(citations=citations, evidence_terms=self._extract_keywords(cleaned), target_terms=6),
                citation_count=len(citations),
                drift_flag=False,
            )
            self._record_quality_metric(
                response_type="job_extract",
                grounding_score=self._score_grounding(citations=citations, evidence_terms=self._extract_keywords(cleaned), target_terms=6),
                citation_count=len(citations),
                confidence=confidence,
                drift_flag=False,
                metadata={"source_type": "manual"},
            )
            return {
                "title": "Manual job description",
                "company": "Provided manually",
                "description": cleaned,
                "source": "manual",
                "source_type": "manual",
                "keywords": sorted(self._extract_keywords(cleaned))[:20],
                "citations": citations,
                "confidence": confidence,
            }
        if not url:
            raise ValueError("Either a job description URL or raw text must be provided.")

        source_id = self._validate_job_source(url)
        async with httpx.AsyncClient(timeout=15.0, headers={"User-Agent": "Actypity/2.0"}) as client:
            response = await client.get(url)
            response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        title = soup.title.string.strip() if soup.title and soup.title.string else "Job description"
        description = self._normalize_whitespace(soup.get_text(" ", strip=True))[:8000]
        citations = self._build_text_citations(
            sources=[(source_id, description)],
            focus_terms=self._extract_keywords(description),
            limit=3,
        )
        confidence = self._confidence_from_grounding(
            grounding_score=self._score_grounding(citations=citations, evidence_terms=self._extract_keywords(description), target_terms=8),
            citation_count=len(citations),
            drift_flag=len(description) < 200,
        )
        self._record_quality_metric(
            response_type="job_extract",
            grounding_score=self._score_grounding(citations=citations, evidence_terms=self._extract_keywords(description), target_terms=8),
            citation_count=len(citations),
            confidence=confidence,
            drift_flag=len(description) < 200,
            metadata={"source_type": source_id, "source": url},
        )
        return {
            "title": title,
            "company": self._guess_company_from_title(title),
            "description": description,
            "source": url,
            "source_type": source_id,
            "keywords": sorted(self._extract_keywords(description))[:20],
            "citations": citations,
            "confidence": confidence,
        }

    def search_jobs(
        self,
        *,
        keywords: List[str],
        locations: List[str],
        sources: List[str],
        created_by: Optional[str],
    ) -> List[Dict[str, Any]]:
        normalized_keywords = [item.strip() for item in keywords if item.strip()]
        normalized_locations = [item.strip() for item in locations if item.strip()] or ["remote"]
        normalized_sources = [item for item in sources if item in TRUSTED_JOB_PORTALS]
        if not normalized_keywords:
            raise ValueError("At least one keyword is required to search jobs.")
        if not normalized_sources:
            normalized_sources = [source["id"] for source in self.trusted_sources()]

        query = " ".join(normalized_keywords)
        results: List[Dict[str, Any]] = []
        for source in normalized_sources:
            config = TRUSTED_JOB_PORTALS[source]
            for location in normalized_locations:
                keyword_token = quote_plus(query)
                location_token = quote_plus(location)
                search_url = config["search_url"].format(
                    keywords=keyword_token,
                    location=location_token,
                )
                results.append(
                    {
                        "id": str(uuid4()),
                        "title": f"{config['label']} search for {query}",
                        "company": config["label"],
                        "location": location,
                        "url": search_url,
                        "source": source,
                        "summary": f"Open a trusted {config['label']} search for {query} in {location}.",
                        "ats_score": None,
                        "description": f"Trusted job search link for {query} in {location} on {config['label']}.",
                        "created_at": datetime.now(timezone.utc).isoformat(),
                    }
                )

        if self.database_client.is_configured:
            self.database_client.save_career_query(
                {
                    "id": str(uuid4()),
                    "query_type": "job_search",
                    "query_text": query,
                    "sources": normalized_sources,
                    "result_count": len(results),
                    "metadata": {"locations": normalized_locations},
                    "created_by": created_by,
                    "created_at": datetime.now(timezone.utc).isoformat(),
                }
            )
            self.database_client.save_job_search_results(results)
        return results

    def evaluate_resume(
        self,
        *,
        resume_text: str,
        jd_text: str,
        model_profile: Optional[str],
    ) -> Dict[str, Any]:
        from agents.resume_skills_agent import ResumeEvaluatorAgent
        agent = ResumeEvaluatorAgent(
            ollama_client=self._ollama_client(),
            llm_client=self.model_router.llm_client if self.model_router.settings.llm_enabled else None,
        )
        result = agent.execute(
            "evaluate resume",
            context={"resume_text": resume_text, "jd_text": jd_text},
        )
        evaluation = result.metadata.get("evaluation", {})
        return {
            "overall_score": evaluation.get("overall_score", 0),
            "grade": evaluation.get("grade", "?"),
            "ats_risk_level": evaluation.get("ats_risk_level", "unknown"),
            "summary": evaluation.get("summary", result.output),
            "dimensions": evaluation.get("dimensions", {}),
            "top_strengths": evaluation.get("top_strengths", []),
            "critical_fixes": evaluation.get("critical_fixes", []),
            "missing_sections": evaluation.get("missing_sections", []),
            "used_llm": result.used_llm,
            "provider": result.metadata.get("provider", ""),
        }

    def write_resume(
        self,
        *,
        resume_text: str,
        jd_text: str,
        target_role: str,
        section: str,
        candidate_name: str,
        model_profile: Optional[str],
    ) -> Dict[str, Any]:
        from agents.resume_skills_agent import ResumeWriterAgent
        agent = ResumeWriterAgent(
            ollama_client=self._ollama_client(),
            llm_client=self.model_router.llm_client if self.model_router.settings.llm_enabled else None,
        )
        result = agent.execute(
            f"write resume for {target_role}",
            context={
                "resume_text": resume_text,
                "jd_text": jd_text,
                "target_role": target_role,
                "section": section,
                "candidate_name": candidate_name,
            },
        )
        written = result.metadata.get("written_content", {})
        return {
            "target_role": target_role,
            "written_content": {
                "professional_summary": written.get("professional_summary", ""),
                "experience_bullets": written.get("experience_bullets", []),
                "skills_section": written.get("skills_section", {}),
                "objective_statement": written.get("objective_statement", ""),
                "keywords_embedded": written.get("keywords_embedded", []),
                "writing_notes": written.get("writing_notes", ""),
            },
            "used_llm": result.used_llm,
            "provider": result.metadata.get("provider", ""),
        }

    def review_resume(
        self,
        *,
        resume_text: str,
        jd_text: str,
        target_role: str,
        model_profile: Optional[str],
    ) -> Dict[str, Any]:
        from agents.resume_skills_agent import ResumeReviewerAgent
        agent = ResumeReviewerAgent(
            ollama_client=self._ollama_client(),
            llm_client=self.model_router.llm_client if self.model_router.settings.llm_enabled else None,
        )
        result = agent.execute(
            "review resume",
            context={
                "resume_text": resume_text,
                "jd_text": jd_text,
                "target_role": target_role,
            },
        )
        review = result.metadata.get("review", {})
        return {
            "overall_verdict": review.get("overall_verdict", result.output),
            "interview_probability": review.get("interview_probability", "medium"),
            "sections": review.get("sections", {}),
            "top_3_immediate_actions": review.get("top_3_immediate_actions", []),
            "red_flags": review.get("red_flags", []),
            "interview_tips": review.get("interview_tips", []),
            "used_llm": result.used_llm,
            "provider": result.metadata.get("provider", ""),
        }

    async def live_hunt_jobs(
        self,
        *,
        resume_text: str,
        location: str,
        experience_years: float,
        model_profile: Optional[str],
    ) -> Dict[str, Any]:
        from .job_search_service import LiveJobSearchService

        resume_keywords = self._extract_keywords(resume_text)
        candidate_name = self._extract_name(resume_text)
        target_roles = self._recommend_roles(resume_keywords, resume_text)[:5]

        searcher = LiveJobSearchService()
        raw_jobs = await searcher.search_for_resume(
            target_roles=target_roles,
            location=location,
            experience_years=experience_years,
            limit_per_role=8,
        )

        analyzed: List[Dict[str, Any]] = []
        for job in raw_jobs:
            fit = self._analyze_job_fit(resume_keywords, job)
            analyzed.append({**job, **fit})

        analyzed.sort(key=lambda j: j["match_score"], reverse=True)
        for j in analyzed:
            score = j["match_score"]
            j["tier"] = "high" if score >= 60 else "medium" if score >= 35 else "stretch"

        high = [j for j in analyzed if j["tier"] == "high"]
        medium = [j for j in analyzed if j["tier"] == "medium"]
        stretch = [j for j in analyzed if j["tier"] == "stretch"]

        return {
            "candidate_name": candidate_name,
            "target_roles": target_roles,
            "total_found": len(analyzed),
            "jobs": analyzed,
            "high_tier": high,
            "medium_tier": medium,
            "stretch_tier": stretch,
        }

    def _extract_name(self, resume_text: str) -> str:
        for line in resume_text.splitlines()[:10]:
            line = line.strip().rstrip('.,|:')
            if re.match(r'^[A-Z][a-z]+(?: [A-Z][a-z]+){1,3}$', line):
                return line
        # fallback: scan first 400 chars for a capitalised 2-word name
        match = re.search(r'\b([A-Z][a-z]{2,}(?:\s+[A-Z][a-z]{2,}){1,2})\b', resume_text[:400])
        return match.group(1) if match else "Candidate"

    def _analyze_job_fit(self, resume_keywords: set[str], job: Dict[str, Any]) -> Dict[str, Any]:
        jd_text = job.get("jd_full") or job.get("jd_snippet", "")
        jd_keywords = self._extract_keywords(jd_text)
        matched = sorted(resume_keywords & jd_keywords)
        missing = sorted(jd_keywords - resume_keywords)
        match_score = self._compute_match_score(matched=len(matched), total=max(len(jd_keywords), 1))

        improvement_areas: List[str] = []
        if missing[:6]:
            improvement_areas.append(f"Add these missing keywords: {', '.join(missing[:6])}")
        if match_score < 60:
            improvement_areas.append("Tailor your professional summary to mirror this role's core requirements")
        if match_score < 40:
            improvement_areas.append("Quantify achievements in domains this JD emphasises")
            improvement_areas.append("Rewrite 2-3 experience bullets to align with the role's language")

        if match_score >= 60:
            ats_summary = f"{match_score}% keyword match — strong fit, apply now."
        elif match_score >= 35:
            ats_summary = f"{match_score}% keyword match — partial fit, tailor your resume before applying."
        else:
            ats_summary = f"{match_score}% keyword match — stretch opportunity, significant gaps to address."

        return {
            "match_score": match_score,
            "matched_keywords": matched[:18],
            "missing_keywords": missing[:12],
            "improvement_areas": improvement_areas,
            "ats_summary": ats_summary,
        }

    def hunt_jobs(
        self,
        *,
        resume_text: str,
        location: str,
        experience_years: float,
        top_count: int,
        model_profile: Optional[str],
    ) -> Dict[str, Any]:
        from agents.job_applicant_agent import JobApplicantAgent
        agent = JobApplicantAgent(
            ollama_client=self._ollama_client(),
            llm_client=self.model_router.llm_client if self.model_router.settings.llm_enabled else None,
        )
        result = agent.execute(
            "job hunt",
            context={
                "resume_text": resume_text,
                "location": location,
                "experience_years": experience_years,
                "top_count": top_count,
            },
        )
        md = result.metadata
        return {
            "candidate_name": md.get("candidate_name", "Candidate"),
            "target_roles": md.get("target_roles", []),
            "total_opportunities": md.get("total", 0),
            "opportunities": md.get("opportunities", []),
            "high_tier": md.get("high_tier", []),
            "medium_tier": md.get("medium_tier", []),
            "stretch_tier": md.get("stretch_tier", []),
            "tailored_applications": md.get("tailored_applications", []),
            "profile": md.get("profile", {}),
        }

    def _ollama_client(self):
        # Lazy-import to avoid circular dependency; container wires the router
        from .local_llm import OllamaClient
        router = self.model_router
        if hasattr(router, "ollama_client") and router.ollama_client and router.ollama_client.enabled:
            return router.ollama_client
        return None

    def analytics(self) -> Dict[str, Any]:
        if self.database_client.is_configured:
            return self.database_client.get_career_analytics()
        return {
            "total_resume_analyses": 0,
            "total_templates": len(DEFAULT_TEMPLATES),
            "total_job_queries": 0,
            "total_job_results": 0,
            "average_match_score": 0.0,
            "top_sources": {},
            "quality_total_evaluations": 0,
            "quality_avg_grounding_score": 0.0,
            "quality_avg_citation_count": 0.0,
            "quality_drift_alerts": 0,
        }

    def _preferred_profile(self, workload: str) -> Optional[str]:
        if workload == "resume" and self.settings.llama_resume_model_path:
            return "llama-local-resume"
        if workload == "job" and self.settings.llama_job_model_path:
            return "llama-local-jd"
        if workload == "template" and self.settings.llama_template_model_path:
            return "llama-local-template"
        if self.settings.llama_model_path:
            return "llama-local-general"
        if self.settings.azure_openai_deployment:
            return "azure-general"
        return None

    def _split_cover_letter_response(self, response: str, target_role: str, company_name: str) -> tuple[str, str, List[str]]:
        lines = [line.strip() for line in response.splitlines() if line.strip()]
        subject_line = next((line.replace("Subject:", "", 1).strip() for line in lines if line.lower().startswith("subject:")), f"Application for {target_role} at {company_name}")
        if "Talking Points:" in response:
            letter_part, talking_part = response.split("Talking Points:", 1)
            letter_lines = [line for line in letter_part.splitlines() if line.strip() and not line.lower().startswith("subject:")]
            talking_points = [
                re.sub(r"^[\-\*\d\.\)\s]+", "", line).strip()
                for line in talking_part.splitlines()
                if re.sub(r"^[\-\*\d\.\)\s]+", "", line).strip()
            ][:5]
            return subject_line, "\n".join(letter_lines).strip(), talking_points
        letter_lines = [line for line in response.splitlines() if line.strip() and not line.lower().startswith("subject:")]
        talking_points = self._extract_bullets(response, limit=3)
        return subject_line, "\n".join(letter_lines).strip(), talking_points

    def _normalize_company_domain(self, company_domain: str) -> str:
        domain = company_domain.strip().lower()
        domain = re.sub(r"^https?://", "", domain)
        domain = domain.split("/")[0]
        return domain.removeprefix("www.")

    def _derive_company_domain(self, url: str) -> str:
        if not url:
            return ""
        host = (urlparse(url).hostname or "").lower()
        return host.removeprefix("www.")

    def _extract_contacts_from_text(self, text: str, *, source: str) -> List[Dict[str, Any]]:
        emails = sorted(set(re.findall(r"[\w\.-]+@[\w\.-]+\.[a-zA-Z]{2,}", text)))
        contacts = []
        for email in emails:
            title = "Recruiting Contact"
            lowered = text.lower()
            if "talent" in lowered or "talent acquisition" in lowered:
                title = "Talent Acquisition"
            elif "human resources" in lowered or "hr" in lowered:
                title = "HR Contact"
            contacts.append(
                {
                    "name": email.split("@", 1)[0].replace(".", " ").replace("_", " ").title(),
                    "title": title,
                    "email": email,
                    "contact_url": None,
                    "source": source,
                    "confidence": "high",
                    "notes": "Discovered from provided text or company page content.",
                }
            )
        return contacts

    def _company_lookup_pages(self, company_domain: str) -> List[str]:
        if not company_domain:
            return []
        return [
            f"https://{company_domain}",
            f"https://{company_domain}/careers",
            f"https://{company_domain}/jobs",
            f"https://{company_domain}/contact",
            f"https://{company_domain}/about",
            f"https://{company_domain}/team",
        ]

    def _build_contact_lookup_urls(self, company_name: str, company_domain: str, target_role: str) -> List[str]:
        urls = []
        keyword = quote_plus(f"{company_name} recruiter {target_role}".strip())
        if company_domain:
            urls.extend(self._company_lookup_pages(company_domain))
        urls.extend(
            [
                f"https://www.linkedin.com/search/results/people/?keywords={keyword}",
                f"https://www.linkedin.com/company/{quote_plus(company_name.lower().replace(' ', '-'))}/people/",
            ]
        )
        return urls

    def _inferred_role_mailboxes(self, company_name: str, company_domain: str) -> List[Dict[str, Any]]:
        aliases = ["careers", "recruiting", "recruiter", "talent", "talent.acquisition", "hr", "jobs"]
        contacts = []
        for alias in aliases:
            contacts.append(
                {
                    "name": f"{company_name} Recruiting",
                    "title": "Recruiting Mailbox",
                    "email": f"{alias}@{company_domain}",
                    "contact_url": f"https://{company_domain}/careers",
                    "source": "inferred-company-pattern",
                    "confidence": "low",
                    "notes": "Inferred from common recruiting mailbox patterns. Verify before outreach.",
                }
            )
        return contacts

    def _dedupe_contacts(self, contacts: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        deduped: Dict[str, Dict[str, Any]] = {}
        for contact in contacts:
            key = (contact.get("email") or contact.get("contact_url") or contact.get("name") or str(uuid4())).lower()
            if key not in deduped:
                deduped[key] = contact
        return list(deduped.values())[:20]

    def _record_quality_metric(
        self,
        *,
        response_type: str,
        grounding_score: int,
        citation_count: int,
        confidence: str,
        drift_flag: bool,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> None:
        if not self.database_client.is_configured:
            return
        try:
            self.database_client.record_response_quality_metric(
                response_type=response_type,
                grounding_score=grounding_score,
                citation_count=citation_count,
                confidence=confidence,
                drift_flag=drift_flag,
                metadata=metadata,
            )
        except Exception:
            pass

    def _shared_evidence_terms(self, resume_text: str, jd_text: str, limit: int = 6) -> List[str]:
        resume_keywords = self._extract_keywords(resume_text)
        jd_keywords = self._extract_keywords(jd_text)
        return sorted(resume_keywords & jd_keywords)[:limit]

    def _build_text_citations(
        self,
        *,
        sources: List[tuple[str, str]],
        focus_terms: set[str] | List[str],
        limit: int = 3,
    ) -> List[Dict[str, Any]]:
        terms = [term for term in focus_terms if term][:8]
        citations: List[Dict[str, Any]] = []
        seen: set[tuple[str, str]] = set()
        for source_name, raw_text in sources:
            if not raw_text:
                continue
            normalized = self._normalize_whitespace(raw_text)
            sentences = re.split(r"(?<=[\.\!\?])\s+", normalized)
            for sentence in sentences:
                sentence_clean = self._normalize_whitespace(sentence)
                if len(sentence_clean) < 30:
                    continue
                matched_terms = [term for term in terms if term.lower() in sentence_clean.lower()]
                if not matched_terms:
                    continue
                excerpt = sentence_clean[:280]
                key = (source_name, excerpt)
                if key in seen:
                    continue
                seen.add(key)
                citations.append(
                    {
                        "source_type": source_name,
                        "excerpt": excerpt,
                        "matched_terms": matched_terms[:4],
                    }
                )
                if len(citations) >= limit:
                    return citations
        return citations

    def _score_grounding(
        self,
        *,
        citations: List[Dict[str, Any]],
        evidence_terms: set[str] | List[str],
        target_terms: int,
    ) -> int:
        term_count = min(len(list(evidence_terms)), target_terms)
        citation_score = min(len(citations) * 25, 75)
        term_score = min(term_count * 5, 25)
        return max(0, min(100, citation_score + term_score))

    def _score_resume_chat_grounding(self, contexts: List[Dict[str, Any]]) -> int:
        avg_similarity = 0.0
        if contexts:
            avg_similarity = sum(float(item.get("score", 0.0)) for item in contexts[:3]) / min(len(contexts), 3)
        normalized_similarity = max(0, min(25, round(avg_similarity * 25)))
        return max(0, min(100, len(contexts[:3]) * 25 + normalized_similarity))

    def _confidence_from_grounding(
        self,
        *,
        grounding_score: int,
        citation_count: int,
        drift_flag: bool,
    ) -> str:
        if drift_flag and grounding_score < 50:
            return "low"
        if grounding_score >= 75 and citation_count >= 2:
            return "high"
        if grounding_score >= 45 and citation_count >= 1:
            return "medium"
        return "low"

    def _contact_grounding_score(self, *, verified_count: int, total_count: int) -> int:
        if total_count <= 0:
            return 0
        ratio = verified_count / total_count
        return max(0, min(100, round(ratio * 100)))

    def _contact_provenance(self, contacts: List[Dict[str, Any]], lookup_urls: List[str]) -> List[str]:
        provenance = {str(contact.get("source")) for contact in contacts if contact.get("source")}
        provenance.update(url for url in lookup_urls[:3] if url)
        return sorted(provenance)[:8]

    def _validate_job_source(self, url: str) -> str:
        hostname = (urlparse(url).hostname or "").lower()
        for source_id, config in TRUSTED_JOB_PORTALS.items():
            if hostname in config["hosts"]:
                return source_id
        raise ValueError("Only trusted job portal URLs are accepted.")

    def _extract_keywords(self, text: str) -> set[str]:
        # Only match alphanumeric + technical chars (+#); no dots to avoid URL/company-name noise
        tokens = re.findall(r"[A-Za-z][A-Za-z0-9+#]{2,}", text.lower())
        stop_words = {
            "with", "from", "that", "this", "have", "will", "your", "their", "about",
            "role", "team", "work", "experience", "years", "skills", "using", "build",
            "strong", "ability", "preferred", "required", "responsible", "you", "the",
            "and", "for", "are", "our", "can", "has", "not", "all", "but", "its",
            "who", "how", "was", "did", "new", "also", "any", "get", "one", "may",
            "job", "include", "including", "position", "company", "candidates",
            "opportunity", "looking", "seeking", "join", "part", "full", "time",
            "remote", "hybrid", "onsite", "apply", "email", "send", "resume", "please",
        }
        return {token for token in tokens if token not in stop_words and len(token) > 2}

    def _extract_bullets(self, text: str, limit: int) -> List[str]:
        lines = [
            self._normalize_whitespace(line)
            for line in text.splitlines()
            if self._normalize_whitespace(line)
        ]
        return lines[:limit]

    def _build_resume_suggestions(self, *, matched: List[str], missing: List[str], strengths: List[str]) -> List[str]:
        suggestions: List[str] = []
        if missing:
            suggestions.append(f"Add evidence for missing job keywords: {', '.join(missing[:6])}.")
        if strengths:
            suggestions.append(f"Rewrite top bullets to foreground impact: {strengths[0][:120]}.")
        if not matched:
            suggestions.append("Mirror more role-specific terminology from the job description in the summary and experience sections.")
        suggestions.append("Quantify outcomes with metrics, scale, ownership, or business impact wherever possible.")
        return suggestions[:4]

    def _fix_pdf_text(self, text: str) -> str:
        # Split camelCase run-ons produced by some PDF extractors (e.g. "abilityto" → no fix, but "leadershipAI" → "leadership AI")
        text = re.sub(r'([a-z]{2})([A-Z][a-z])', r'\1 \2', text)
        # Re-space lowercase run-ons where a known boundary word is fused (e.g. "aboutusing" → "about using")
        boundary_words = (
            "about|ability|across|also|and|are|as|at|be|been|being|between|by|can|career|"
            "with|within|without|worked|working|years|using|used|through|to|the|that|their|"
            "of|on|or|in|into|is|it|for|from|has|have|helping|including|leveraging|managing|"
            "multiple|new|not|over|product|providing|responsible|since|skills|such|team|led"
        )
        text = re.sub(rf'({boundary_words})([A-Z])', r'\1 \2', text, flags=re.IGNORECASE)
        # Fix missing space before capitalised words that immediately follow lowercase (heurisitc for missing spaces)
        text = re.sub(r'([a-z]{3,})([A-Z][a-z]{2,})', r'\1 \2', text)
        # Collapse multiple spaces
        text = re.sub(r' {2,}', ' ', text)
        return text

    def _detect_seniority(self, keywords: set[str], resume_text: str) -> str:
        text_lower = resume_text.lower()
        vp_signals = {"vp", "vice president", "executive vice", "evp", "svp", "chief", "cto", "cdo", "cpo"}
        director_signals = {"director", "associate director", "senior director", "head of"}
        manager_signals = {"manager", "lead", "principal", "staff", "senior"}
        if vp_signals & keywords or any(s in text_lower for s in vp_signals):
            return "vp"
        if director_signals & keywords or any(s in text_lower for s in director_signals):
            return "director"
        if manager_signals & keywords:
            return "manager"
        # Heuristic: count experience year mentions
        year_matches = re.findall(r'(\d+)\s*\+?\s*years?', text_lower)
        max_years = max((int(y) for y in year_matches), default=0)
        if max_years >= 15:
            return "vp"
        if max_years >= 8:
            return "director"
        if max_years >= 4:
            return "manager"
        return "ic"

    def _recommend_roles(self, keywords: set[str], resume_text: str = "") -> List[str]:
        seniority = self._detect_seniority(keywords, resume_text)
        is_ai = bool({"ai", "genai", "llm", "langchain", "rag", "mlops", "pytorch", "tensorflow",
                      "machine", "learning", "nlp", "data", "ml"} & keywords)
        is_data = bool({"data", "analytics", "bi", "sql", "tableau", "powerbi", "databricks"} & keywords)
        is_frontend = bool({"react", "typescript", "frontend", "angular", "vue"} & keywords)
        is_backend = bool({"python", "fastapi", "backend", "django", "flask", "node"} & keywords)
        is_consulting = bool({"consulting", "strategy", "engagement", "advisory", "mckinsey",
                              "deloitte", "kpmg", "pwc", "accenture"} & keywords)

        if seniority == "vp":
            if is_ai:
                return ["VP AI / Head of AI", "Chief AI Officer", "SVP Data & AI", "Director of AI Labs", "VP Engineering (AI)"]
            if is_data:
                return ["VP Data & Analytics", "Chief Data Officer", "Head of Data Science"]
            if is_consulting:
                return ["Partner / Principal", "VP Strategy & AI", "Managing Director – Digital"]
            return ["VP Engineering", "SVP Technology", "CTO – Scale-up"]
        if seniority == "director":
            if is_ai:
                return ["Director of AI", "Head of GenAI", "Director ML Engineering", "Principal AI Architect"]
            if is_data:
                return ["Director Data Science", "Head of Analytics", "Principal Data Engineer"]
            if is_frontend:
                return ["Director of Engineering – Frontend", "Head of Product Engineering"]
            if is_backend:
                return ["Director of Backend Engineering", "Principal Platform Engineer"]
            return ["Director of Engineering", "Head of Technology", "Senior Engineering Manager"]
        if seniority == "manager":
            if is_ai:
                return ["Senior ML Engineer", "AI/ML Tech Lead", "Senior Data Scientist", "ML Platform Lead"]
            if is_frontend:
                return ["Senior Frontend Engineer", "Tech Lead – Frontend", "Staff Product Engineer"]
            if is_backend:
                return ["Senior Backend Engineer", "Platform Engineer", "Applied AI Engineer"]
            return ["Engineering Manager", "Tech Lead", "Senior Software Engineer"]
        # IC / fresher
        if is_ai:
            return ["Machine Learning Engineer", "Data Scientist", "AI Engineer", "NLP Engineer"]
        if {"react", "typescript", "frontend"} & keywords:
            return ["Frontend Engineer", "React Developer", "UI Engineer"]
        if {"python", "fastapi", "backend"} & keywords:
            return ["Backend Engineer", "Software Engineer", "SDE-1"]
        return ["Software Engineer", "Graduate Engineer Trainee", "Associate Software Developer"]

    def _compute_match_score(self, *, matched: int, total: int) -> int:
        return max(0, min(100, round((matched / total) * 100)))

    def _template_sections_for_role(self, target_role: str) -> List[str]:
        role = target_role.lower()
        if "design" in role:
            return ["Header", "Profile", "Case Studies", "Experience", "Toolbox", "Education"]
        if "engineer" in role or "developer" in role:
            return ["Header", "Summary", "Technical Skills", "Experience", "Projects", "Education"]
        return ["Header", "Summary", "Experience", "Achievements", "Skills", "Education"]

    def _design_tokens_for_style(self, style: str) -> Dict[str, str]:
        style_key = style.lower()
        if "minimal" in style_key:
            return {
                "primary": "#102a43",
                "accent": "#f0b429",
                "surface": "#f7fafc",
                "font_heading": "Manrope",
                "font_body": "Source Sans 3",
            }
        if "creative" in style_key or "bold" in style_key:
            return {
                "primary": "#2b2d42",
                "accent": "#ef476f",
                "surface": "#fff8f0",
                "font_heading": "Space Grotesk",
                "font_body": "Plus Jakarta Sans",
            }
        return {
            "primary": "#0b132b",
            "accent": "#3a86ff",
            "surface": "#f8f9fb",
            "font_heading": "Sora",
            "font_body": "Inter",
        }

    def _preview_markdown(self, name: str, target_role: str, sections: List[str]) -> str:
        return (
            f"# {name}\n\n"
            f"Designed for **{target_role}**.\n\n"
            f"Sections:\n- " + "\n- ".join(sections)
        )

    def _guess_company_from_title(self, title: str) -> str:
        if " at " in title.lower():
            return title.split(" at ", 1)[-1].strip()
        return "Unknown"

    def _normalize_whitespace(self, value: str) -> str:
        return re.sub(r"\s+", " ", value).strip()
